# src/io/loaders.py
"""Enhanced document loaders for PDF, Word, and Excel files."""

import logging
from pathlib import Path
from typing import List, Tuple, Optional

from pypdf import PdfReader
from docx import Document
from openpyxl import load_workbook

logger = logging.getLogger(__name__)


def pdf_to_pages(path: str) -> List[Tuple[int, str]]:
    """Extract text from PDF, one tuple per page."""
    try:
        reader = PdfReader(path)
        pages = []
        for i, page in enumerate(reader.pages, start=1):
            text = page.extract_text() or ""
            pages.append((i, text))
        logger.info(f"Extracted {len(pages)} pages from PDF: {Path(path).name}")
        return pages
    except Exception as e:
        logger.error(f"Failed to load PDF {path}: {e}")
        raise


def docx_to_pages(path: str) -> List[Tuple[int, str]]:
    """
    Extract text from Word document with structure preservation.
    
    Strategy:
    - Each major heading (Heading 1, Heading 2) starts a new "page"
    - Accumulate paragraphs until next heading
    - Preserve tables as formatted text
    - Return list of (section_num, text) tuples
    """
    try:
        doc = Document(path)
        sections: List[Tuple[int, str]] = []
        current_section_num = 1
        current_section_text: List[str] = []
        current_heading = "Introduction"
        
        for element in doc.element.body:
            # Handle paragraphs
            if element.tag.endswith('p'):
                para = None
                for p in doc.paragraphs:
                    if p._element == element:
                        para = p
                        break
                
                if para is None:
                    continue
                
                # Check if this is a heading
                if para.style.name.startswith('Heading'):
                    # Save previous section if it has content
                    if current_section_text:
                        section_text = f"[Section: {current_heading}]\n\n" + "\n".join(current_section_text)
                        sections.append((current_section_num, section_text))
                        current_section_num += 1
                        current_section_text = []
                    
                    # Start new section
                    current_heading = para.text.strip() or f"Section {current_section_num}"
                
                # Add paragraph text (including headings as context)
                if para.text.strip():
                    current_section_text.append(para.text)
            
            # Handle tables
            elif element.tag.endswith('tbl'):
                table = None
                for t in doc.tables:
                    if t._element == element:
                        table = t
                        break
                
                if table is None:
                    continue
                
                # Convert table to readable text format
                table_text = _format_table(table)
                if table_text:
                    current_section_text.append(f"\n[TABLE]\n{table_text}\n[/TABLE]\n")
        
        # Add final section
        if current_section_text:
            section_text = f"[Section: {current_heading}]\n\n" + "\n".join(current_section_text)
            sections.append((current_section_num, section_text))
        
        # Fallback: if no sections were created, treat entire doc as one page
        if not sections:
            full_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            sections.append((1, full_text))
        
        logger.info(f"Extracted {len(sections)} sections from Word doc: {Path(path).name}")
        return sections
    
    except Exception as e:
        logger.error(f"Failed to load Word doc {path}: {e}")
        raise


def _format_table(table) -> str:
    """Convert Word table to plain text with column alignment."""
    try:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):  # Skip empty rows
                rows.append(" | ".join(cells))
        return "\n".join(rows) if rows else ""
    except Exception as e:
        logger.warning(f"Failed to format table: {e}")
        return "[Table extraction failed]"


def excel_to_pages(path: str) -> List[Tuple[int, str]]:
    """
    Extract text from Excel file.
    
    Strategy:
    - Each sheet becomes one "page"
    - Extract all non-empty cells with row/column context
    - Preserve table structure where possible
    - Include sheet name as section header
    """
    try:
        workbook = load_workbook(path, data_only=True)
        sheets: List[Tuple[int, str]] = []
        
        for sheet_num, sheet_name in enumerate(workbook.sheetnames, start=1):
            sheet = workbook[sheet_name]
            
            # Get the used range to avoid empty cells
            if sheet.max_row == 1 and sheet.max_column == 1:
                # Skip empty sheets
                continue
            
            sheet_text_parts = [f"[Sheet: {sheet_name}]", ""]
            
            # Try to detect if this is a table with headers
            first_row = list(sheet.iter_rows(min_row=1, max_row=1, values_only=True))[0]
            has_headers = any(first_row) and all(
                isinstance(cell, str) or cell is None 
                for cell in first_row
            )
            
            if has_headers:
                # Extract as table with headers
                headers = [str(cell) if cell is not None else "" for cell in first_row]
                sheet_text_parts.append("| " + " | ".join(headers) + " |")
                sheet_text_parts.append("|" + "|".join(["---"] * len(headers)) + "|")
                
                for row in sheet.iter_rows(min_row=2, values_only=True):
                    row_values = [str(cell) if cell is not None else "" for cell in row]
                    if any(row_values):  # Skip empty rows
                        sheet_text_parts.append("| " + " | ".join(row_values) + " |")
            else:
                # Extract as unstructured text (cell by cell)
                for row_idx, row in enumerate(sheet.iter_rows(values_only=True), start=1):
                    row_values = []
                    for col_idx, cell in enumerate(row, start=1):
                        if cell is not None and str(cell).strip():
                            row_values.append(f"[R{row_idx}C{col_idx}] {cell}")
                    
                    if row_values:
                        sheet_text_parts.append(" | ".join(row_values))
            
            sheet_text = "\n".join(sheet_text_parts)
            
            if sheet_text.strip():
                sheets.append((sheet_num, sheet_text))
        
        # Fallback: if no sheets extracted, return empty
        if not sheets:
            logger.warning(f"No content extracted from Excel file: {Path(path).name}")
            sheets.append((1, f"[Empty Excel file: {Path(path).name}]"))
        
        logger.info(f"Extracted {len(sheets)} sheets from Excel: {Path(path).name}")
        return sheets
    
    except Exception as e:
        logger.error(f"Failed to load Excel file {path}: {e}")
        raise


def load_document(path: str) -> List[Tuple[int, str]]:
    """
    Universal document loader - auto-detects format and routes to appropriate loader.
    
    Returns:
        List of (page_num/section_num, text) tuples
    """
    file_path = Path(path)
    suffix = file_path.suffix.lower()
    
    loaders = {
        '.pdf': pdf_to_pages,
        '.docx': docx_to_pages,
        '.doc': docx_to_pages,  # python-docx handles .doc too (with limitations)
        '.xlsx': excel_to_pages,
        '.xls': excel_to_pages,
    }
    
    loader = loaders.get(suffix)
    if loader is None:
        raise ValueError(
            f"Unsupported file type: {suffix}. "
            f"Supported types: {', '.join(loaders.keys())}"
        )
    
    logger.info(f"Loading {suffix} file: {file_path.name}")
    return loader(str(path))